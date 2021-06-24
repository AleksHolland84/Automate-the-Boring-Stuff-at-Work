# Elevplan_copy.py Copies a base-file and save one file 
# for each member in a list

import shutil, os, re, datetime
from docx import Document

# program variables and meta
folder = ''
documentnavn = 'Elevplan'
base_file = 'skabelon.docx'
subfolder = 'Elevplaner'
klasse = input('Hvilken klasse?: ')
elever = 'elev liste.txt'
f_type = '.txt'


path = '' # will contain path to subfolder

def dater():
    now = datetime.date.today()
    return now

def create_folder():
    global folder, subfolder, path
    folder = os.path.abspath(folder) # make sure path to folder is absolute
    new_folder_path = folder + '\\' + (subfolder + '_' + klasse) 
    if not os.path.exists(new_folder_path):
        os.mkdir(new_folder_path)
        path = new_folder_path
    else:
        path = new_folder_path

def create_list(filename, f_type):
    filetype = re.compile(f'{f_type}$')
    # takes a file as an argument and returns a list
    # contaning all the names in the file
    names = []
    filename = input("Filnavn på liste af elever?: ")
    mo = filetype.search(filename)
    try:
        if filename == '':
            filename = elever 
        elif mo:
            # if .txt in filename
            pass
        else: 
            filename = filename + '.txt'
        with open(filename, 'r', encoding='utf-8') as file:
            for name in file:
                names.append(name.strip(' \n\t'))
        return names
    except FileNotFoundError as fnf_error:
        print(fnf_error)
        print(f"There's no file named: {filename}")
    except AttributeError as e:
        print(e)
        print('program error with regex at createlist() ')
    

def edit_header(file_name, name):
    try:
        doc = Document(file_name)
        doc.tables #a list of all tables in document
        doc.tables[0].cell(1,1).text = name
        doc.tables[0].cell(2,1).text = klasse
        doc.tables[0].cell(3,1).text = str((dater().year)-1) + '-' + str(dater().year) 
        doc.save(file_name)

    except docx.opc.exceptions.PackageNotFoundError as docx_error:
        print(docx_error)
        print(f'*** ERROR at edit_header() - {file_name} not found or is already opened***')


def copy_frame(folder, base_file, new_path, name):
    # copies a base_file and saves it as documentnavn + year + name of the name-argument
    foldername = os.path.abspath(new_path)
    new_file_path = foldername + f'\\{documentnavn} {dater().year} {name}.docx'
    try:
        if not os.path.exists(new_file_path): # check if file exists
            shutil.copy(os.path.join(folder, base_file), new_file_path)
            edit_header(new_file_path, name)
            print(f'Created {os.path.basename(new_file_path)}')
        else:
            print(f"{os.path.basename(new_file_path)} ALREADY EXISTS")
    except FileNotFoundError as fnf_error:
        print(fnf_error)
        print('*** ERROR at copy_frame() ***')





if __name__ == '__main__':
    elever = create_list(elever, f_type)

    create_folder()

    for elev in elever:
        copy_frame(folder, base_file, path, elev)
