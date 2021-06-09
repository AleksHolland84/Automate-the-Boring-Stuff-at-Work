#!C:\Users\alek1024\OneDrive - Viborg Skoler\2020-2021 Lærermappe\Elevplaner_\.Elevplan_Copy@3.9.2\Scripts\python.exe
# elevplan_copy.py Copies a base-file and save one file 
# for each member in a list

import shutil, os, re, datetime
from docx import Document

# program variables and meta
folder = ''
documentnavn = 'Elevplan'
base_file = 'skabelon.docx'
subfolder = 'Elevplaner'
klasse = input('Hvilken klasse?: ')
elever = klasse + '.txt'


path = '' # will contain path to subfolder

def dater():
    now = datetime.date.today()
    return now

def create_folder():
    global folder, subfolder, path
    folder = os.path.abspath(folder) # make sure path to folder is absolute
    new_folder_path = folder + '\\' + (subfolder + '_' + klasse) 
    if not os.path.exists(new_folder_path):
        os.mkdir(new_folder_path)
        path = new_folder_path
    else:
        path = new_folder_path

def create_list(filename):
    # takes a file as an argument and returns a list
    # contaning all the names in the file
    names = []
    filename = input("Filnavn på liste af elever?: ")
    filename = filename + '.txt'
    if filename == '':
        filename = elever
        with open(filename, 'r', encoding='utf-8') as file:
            for name in file:
                names.append(name.strip(' \n\t'))
        return names    
    else:
        with open(filename, 'r', encoding='utf-8') as file:
            for name in file:
                names.append(name.strip(' \n\t'))
        return names


def edit_header(file_name, name):
    doc = Document(file_name)
    doc.tables #a list of all tables in document
    if doc.tables[0].cell(1,1).text == '':
        doc.tables[0].cell(1,1).text = name

        if doc.tables[0].cell(2,1).text == '':
            doc.tables[0].cell(2,1).text = klasse

            if doc.tables[0].cell(3,1).text == '':
                doc.tables[0].cell(3,1).text = str((dater().year)-1) + '-' + str(dater().year) 
        doc.save(file_name)
    else:
        print('Cell not empty!')


def copy_frame(folder, base_file, new_path, name):
    # copies a base_file and saves it as elevplan + name of the name-argument
    foldername = os.path.abspath(new_path)
    new_file_path = foldername + f'\\{documentnavn} {dater().year} {name}.docx'
    try:
        if not os.path.exists(new_file_path): # check if file exists
            shutil.copy(os.path.join(folder, base_file), new_file_path)
            edit_header(new_file_path, name)
        else:
            print(f"{os.path.basename(new_file_path)} ALLREADY EXISTS")
    except FileNotFoundError as fnf_error:
        print(fnf_error)

def edit_header_tester(file_name, name):
    doc = Document(file_name)
    doc.tables #a list of all tables in document

    for row in doc.tables[0].rows:
        for cell in row.cells:
            if cell.text == 'Elev:':
                print(cell)
                print('test')
            else:
                print(cell.text)



if __name__ == '__main__':
    elever = create_list(elever)

    create_folder()

    for elev in elever:
        copy_frame(folder, base_file, path, elev)
